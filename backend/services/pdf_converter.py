import os
import io
import re
import pypdfium2 as pdfium
import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from paddleocr import PaddleOCR

class PDFConverterService:
    def __init__(self):
        # Initialize PaddleOCR for Arabic
        self.ocr = PaddleOCR(use_angle_cls=True, lang='ar')

    def process_arabic_text(self, text: str) -> str:
        """
        Processes Arabic text to ensure characters are connected correctly
        and the Right-to-Left (RTL) reading order is preserved.
        """
        # 1. PaddleOCR returns Arabic in Visual Order (Left-to-Right).
        # We must reverse the entire string to get Logical Order.
        reversed_text = text[::-1]
        
        # 2. Reversing the string messed up English numbers (e.g., "123" became "321").
        # We need to find all numbers/English characters and reverse them back to LTR.
        def reverse_match(match):
            return match.group(0)[::-1]
            
        pattern = r'[0-9A-Za-z\.,:\-\%]+(?:\s+[0-9A-Za-z\.,:\-\%]+)*'
        logical_text = re.sub(pattern, reverse_match, reversed_text)
        
        # 3. Swap brackets so MS Word's native Bidi renders them correctly
        logical_text = logical_text.replace('(', 'TEMP_B').replace(')', '(').replace('TEMP_B', ')')
        logical_text = logical_text.replace('[', 'TEMP_B').replace(']', '[').replace('TEMP_B', ']')
        logical_text = logical_text.replace('{', 'TEMP_B').replace('}', '{').replace('TEMP_B', '}')
        logical_text = logical_text.replace('«', 'TEMP_B').replace('»', '«').replace('TEMP_B', '»')
        logical_text = logical_text.replace('﴾', 'TEMP_B').replace('﴿', '﴾').replace('TEMP_B', '﴿')
        
        # 4. Reshape the text to connect Arabic letters correctly
        reshaped_text = arabic_reshaper.reshape(logical_text)
        
        # 5. Fix the Bidi direction (so words flow RTL correctly)
        bidi_text = get_display(reshaped_text)
        
        return bidi_text

    def convert_pdf_to_docx(self, pdf_path: str, docx_path: str):
        """
        Converts an Arabic PDF file to a Word Document (.docx) using AI OCR.
        Preserves Arabic text formatting and readability.
        """
        pdf_document = pdfium.PdfDocument(pdf_path)
        doc = Document()
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            bitmap = page.render(scale=300/72)
            pil_image = bitmap.to_pil()
            
            temp_image_path = f"{pdf_path}_page_{page_num}.png"
            pil_image.save(temp_image_path)
            
            try:
                result = self.ocr.ocr(temp_image_path)
                
                if result and result[0]:
                    # Group bounding boxes into lines based on Y coordinate
                    lines = []
                    current_line = []
                    last_y = None
                    
                    # Sort boxes primarily by Y (top-to-bottom)
                    boxes = sorted(result[0], key=lambda b: b[0][0][1])
                    
                    for box in boxes:
                        coords = box[0]
                        text = box[1][0]
                        
                        center_y = sum(p[1] for p in coords) / 4
                        center_x = sum(p[0] for p in coords) / 4
                        
                        if last_y is None or abs(center_y - last_y) > 15:
                            if current_line:
                                lines.append(current_line)
                            current_line = [(center_x, text)]
                            last_y = center_y
                        else:
                            current_line.append((center_x, text))
                            last_y = (last_y + center_y) / 2
                            
                    if current_line:
                        lines.append(current_line)
                        
                    for line_boxes in lines:
                        # Sort words in the line by X coordinate LTR
                        line_boxes = sorted(line_boxes, key=lambda b: b[0], reverse=False)
                        combined_visual_text = " ".join([b[1] for b in line_boxes])
                        
                        # Apply Arabic reshaping and Bidi algorithm
                        final_text = self.process_arabic_text(combined_visual_text)
                        
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        
                        p_pr = p._element.get_or_add_pPr()
                        bidi_p = OxmlElement('w:bidi')
                        p_pr.append(bidi_p)
                        
                        run = p.add_run(final_text)
                        run.font.rtl = True
                
                if page_num < len(pdf_document) - 1:
                    doc.add_page_break()
                    
            finally:
                if os.path.exists(temp_image_path):
                    os.remove(temp_image_path)
        
        doc.save(docx_path)
        pdf_document.close()
